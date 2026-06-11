"""DOCX text extraction: paragraphs + tables, in document order where possible."""
import docx as python_docx


def extract(filepath: str) -> str:
    document = python_docx.Document(filepath)
    parts: list[str] = []

    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    # Tables rendered as pipe-separated rows so the LLM can read them
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
