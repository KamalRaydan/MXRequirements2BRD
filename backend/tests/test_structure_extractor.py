"""StructureExtractor: heading hierarchy order + canonical id mapping (spec §17.1)."""
import docx as python_docx
import pytest

from services import structure_extractor


def _make_docx(path, headings):
    """headings: list of (title, level). Level 0 = plain body paragraph."""
    document = python_docx.Document()
    for title, level in headings:
        if level == 0:
            document.add_paragraph(title)
        else:
            document.add_heading(title, level=level)
    document.save(str(path))
    return str(path)


def test_headings_in_document_order_with_levels(tmp_path):
    path = _make_docx(tmp_path / "branded.docx", [
        ("Document Control", 1),
        ("Some body text that is not a heading", 0),
        ("Executive Summary", 1),
        ("Scope", 1),
        ("In Scope", 2),
        ("Out of Scope", 2),
        ("Business Requirements", 1),
        ("Appendix A — Sources", 1),
    ])

    sections = structure_extractor.extract_headings(path)

    assert [s["id"] for s in sections] == [
        "cover", "executive_summary", "scope", "scope_in", "scope_out",
        "requirements", "appendix",
    ]
    assert [s["level"] for s in sections] == [1, 1, 1, 2, 2, 1, 1]
    assert all(s["required"] for s in sections)


def test_unknown_headings_get_slug_ids(tmp_path):
    path = _make_docx(tmp_path / "custom.docx", [
        ("Stakeholder Sign-Off", 1),
        ("Stakeholder Sign-Off", 2),  # duplicate title -> suffixed id
    ])

    sections = structure_extractor.extract_headings(path)

    assert sections[0]["id"] == "stakeholder_sign_off"
    assert sections[1]["id"] == "stakeholder_sign_off_2"
    assert sections[0]["title"] == "Stakeholder Sign-Off"


def test_specific_keywords_win_over_general_ones(tmp_path):
    path = _make_docx(tmp_path / "reqs.docx", [
        ("Non-Functional Requirements", 1),
        ("Integration Requirements", 1),
        ("Business Requirements", 1),
    ])

    sections = structure_extractor.extract_headings(path)

    assert [s["id"] for s in sections] == ["non_functional", "integrations", "requirements"]


def test_headings_deeper_than_3_are_ignored(tmp_path):
    path = _make_docx(tmp_path / "deep.docx", [
        ("Top", 1),
        ("Too deep", 4),
    ])

    sections = structure_extractor.extract_headings(path)

    assert [s["title"] for s in sections] == ["Top"]


def test_document_without_headings_raises(tmp_path):
    path = _make_docx(tmp_path / "plain.docx", [("Just a paragraph", 0)])

    with pytest.raises(ValueError):
        structure_extractor.extract_headings(path)
