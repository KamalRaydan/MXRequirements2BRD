"""Branded template upload/preview/removal routes (spec §10.2, Milestone 2)."""
import io

import docx as python_docx
import pytest
from fastapi.testclient import TestClient

import main


@pytest.fixture
def client():
    with TestClient(main.app) as test_client:
        yield test_client


@pytest.fixture
def project(client, tmp_path):
    return client.post("/projects", json={
        "client_name": "Brand Co", "project_name": "Branding Test",
        "project_date": "2026-06-12", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "brand"),
    }).json()


def _docx_bytes(headings):
    document = python_docx.Document()
    for title, level in headings:
        document.add_heading(title, level=level)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_upload_branding_returns_heading_preview(client, project):
    payload = _docx_bytes([("Executive Summary", 1), ("Business Requirements", 1)])
    response = client.put(
        f"/projects/{project['id']}/branding",
        files={"file": ("client-template.docx", payload)},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["branded_docx_path"].endswith("branding/reference.docx")
    assert [h["id"] for h in body["headings"]] == ["executive_summary", "requirements"]

    # Preview is also available on a fresh GET (page reload)
    preview = client.get(f"/projects/{project['id']}/branding").json()
    assert len(preview["headings"]) == 2


def test_upload_branding_returns_and_persists_profile(client, project):
    """The visual branding profile is returned on upload and on later GETs (M6)."""
    response = client.put(
        f"/projects/{project['id']}/branding",
        files={"file": ("t.docx", _docx_bytes([("Scope", 1)]))},
    )
    profile = response.json()["profile"]
    assert profile is not None
    # A heading-only template has no logo and uses theme fonts
    assert profile["logo_found"] is False
    assert "table_style" in profile

    # Profile survives a page reload (loaded from profile.json)
    assert client.get(f"/projects/{project['id']}/branding").json()["profile"] == profile


def test_upload_branding_rejects_non_docx(client, project):
    response = client.put(
        f"/projects/{project['id']}/branding",
        files={"file": ("notes.txt", b"not a template")},
    )
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "INVALID_TEMPLATE"


def test_upload_branding_rejects_docx_without_headings(client, project):
    document = python_docx.Document()
    document.add_paragraph("No headings here")
    buffer = io.BytesIO()
    document.save(buffer)

    response = client.put(
        f"/projects/{project['id']}/branding",
        files={"file": ("flat.docx", buffer.getvalue())},
    )
    assert response.status_code == 422
    # A rejected template must not be left active on the project
    assert client.get(f"/projects/{project['id']}/branding").json()["branded_docx_path"] is None


def test_delete_branding_reverts_to_default(client, project):
    client.put(
        f"/projects/{project['id']}/branding",
        files={"file": ("t.docx", _docx_bytes([("Scope", 1)]))},
    )

    assert client.delete(f"/projects/{project['id']}/branding").status_code == 204
    preview = client.get(f"/projects/{project['id']}/branding").json()
    assert preview == {"branded_docx_path": None, "headings": [], "profile": None}
