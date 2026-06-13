"""BrandingProfile extraction from a reference DOCX (spec §18, Milestone 6)."""
import struct
import zlib
from io import BytesIO

import docx as python_docx
from docx.shared import Inches, Pt, RGBColor

from services import branding_profile


def _png_1x1() -> bytes:
    """A real 1x1 RGB PNG with correct chunk CRCs — python-docx's image parser is
    strict, and Pillow isn't a dependency, so we hand-build a minimal valid file."""
    def chunk(kind: bytes, data: bytes) -> bytes:
        body = kind + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)   # 1x1, 8-bit, RGB
    idat = zlib.compress(b"\x00\x00\x00\x00")              # filter byte + one black pixel
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


_PNG_1X1 = _png_1x1()


def _branded_docx(path, *, with_logo=True):
    """A reference DOCX with a known body/heading font, heading colour, a styled
    table, and (optionally) a logo image in the page header."""
    document = python_docx.Document()
    document.styles["Normal"].font.name = "Georgia"
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Heading 1"].font.name = "Georgia"
    document.styles["Heading 1"].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    document.add_heading("Executive Summary", level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    if with_logo:
        run = document.sections[0].header.paragraphs[0].add_run()
        run.add_picture(BytesIO(_PNG_1X1), width=Inches(1))

    document.save(str(path))
    return str(path)


def test_extract_profile_reads_fonts_color_table_and_logo(tmp_path):
    reference = _branded_docx(tmp_path / "reference.docx")
    profile = branding_profile.extract_profile(reference, branding_dir=str(tmp_path))

    assert profile["body_font"] == "Georgia"
    assert profile["body_size_pt"] == 11.0
    assert profile["heading_font"] == "Georgia"
    assert profile["heading_color"] == "1F3864"
    assert profile["table_style"] == "Table Grid"
    assert profile["logo_found"] is True
    assert profile["logo_path"] and profile["logo_path"].endswith(".png")
    assert (tmp_path / "logo.png").exists()


def test_extract_profile_without_logo(tmp_path):
    reference = _branded_docx(tmp_path / "reference.docx", with_logo=False)
    profile = branding_profile.extract_profile(reference, branding_dir=str(tmp_path))

    assert profile["logo_found"] is False
    assert profile["logo_path"] is None


def test_save_load_remove_profile_roundtrip(tmp_path):
    reference = _branded_docx(tmp_path / "reference.docx")
    profile = branding_profile.extract_profile(reference, branding_dir=str(tmp_path))

    branding_profile.save_profile(profile, str(tmp_path))
    assert branding_profile.load_profile(str(tmp_path)) == profile

    branding_profile.remove_profile(str(tmp_path))
    assert branding_profile.load_profile(str(tmp_path)) is None
    assert not (tmp_path / "logo.png").exists()
