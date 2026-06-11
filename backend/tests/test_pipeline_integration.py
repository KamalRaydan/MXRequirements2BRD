"""Full pipeline with a mocked LLM: create project -> upload TXT -> wait EXTRACTED
-> generate -> run DONE -> DOCX downloads (spec §17.2)."""
import pytest
from fastapi.testclient import TestClient

import agents.runner
import main
from db.database import SessionLocal
from db.models import PipelineRun
from models.pipeline import (AnalysisDraft, NarrativeSection, NarrativeSet,
                             RequirementDraft)
from services import keystore
from services.progress_bus import bus


class FakeLLM:
    """Stands in for LLMClient — returns canned, schema-valid responses."""

    def __init__(self, api_key, model_id, provider="anthropic"):
        pass

    def complete(self, messages, max_tokens=4096, system=None):
        return "summary text"

    def complete_json(self, messages, schema, max_tokens=8192, system=None):
        if schema is AnalysisDraft:
            return AnalysisDraft(
                requirements=[RequirementDraft(
                    module="WO", title="Auto-close work orders",
                    description="Work orders must auto-close after 30 days",
                    requirement_type="functional", priority="high",
                    source_ref="notes.txt", source_timestamp="2026-05-15T09:00:00Z",
                    sort_order=1,
                )],
                modules_referenced=["WO"],
            )
        return NarrativeSet(narratives=[
            NarrativeSection(section_id="executive_summary",
                             title="Executive Summary", body="The summary."),
        ])


@pytest.fixture
def client(monkeypatch):
    monkeypatch.setattr(agents.runner, "LLMClient", FakeLLM)
    monkeypatch.setattr(keystore, "get_api_key", lambda provider: "test-key")
    with TestClient(main.app) as test_client:
        yield test_client


def test_full_pipeline_with_mock_llm(client, tmp_path):
    project = client.post("/projects", json={
        "client_name": "Test Co", "project_name": "Pipeline Test",
        "project_date": "2026-06-11", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "proj"),
    }).json()

    # Upload — TestClient runs the background extraction before returning control
    upload = client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.txt", b"WOs must auto-close after 30 days.", "text/plain")},
    )
    assert upload.status_code == 201

    sources = client.get(f"/projects/{project['id']}/sources").json()
    assert sources[0]["processing_status"] == "EXTRACTED"

    # Generate — pipeline background task also completes before control returns
    generate = client.post(f"/projects/{project['id']}/generate")
    assert generate.status_code == 202
    run_id = generate.json()["run_id"]

    run = client.get(f"/pipeline/{run_id}").json()
    assert run["status"] == "DONE", run.get("error_message")
    assert run["sources_used_count"] == 1

    download = client.get(f"/pipeline/{run_id}/download")
    assert download.status_code == 200
    assert download.content[:2] == b"PK"  # docx files are zip archives


def test_generate_blocked_without_sources(client, tmp_path):
    project = client.post("/projects", json={
        "client_name": "Empty Co", "project_name": "No Sources",
        "project_date": "2026-06-11", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "empty"),
    }).json()

    response = client.post(f"/projects/{project['id']}/generate")
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "NO_SOURCES"


def test_malformed_pdf_gets_error_status(client, tmp_path):
    """Corrupt-file handling (spec §17.2): a broken PDF must land on ERROR with a
    message, not crash the upload or block the project."""
    project = client.post("/projects", json={
        "client_name": "Bad Co", "project_name": "Corrupt PDF",
        "project_date": "2026-06-11", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "corrupt"),
    }).json()

    upload = client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("broken.pdf", b"this is not a real pdf", "application/pdf")},
    )
    assert upload.status_code == 201

    source = client.get(f"/projects/{project['id']}/sources").json()[0]
    assert source["processing_status"] == "ERROR"
    assert source["error_message"]


def test_cancel_flag_stops_pipeline_between_stages(client, tmp_path):
    """Cancel flow (spec §10.5): with the flag pre-set, the run ends CANCELLED
    and the stream reports stage 'cancelled'."""
    project = client.post("/projects", json={
        "client_name": "Cancel Co", "project_name": "Cancelled Run",
        "project_date": "2026-06-11", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "cancel"),
    }).json()
    client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.txt", b"WOs must auto-close.", "text/plain")},
    )

    # Create the run directly so we can set the cancel flag before the pipeline
    # starts (TestClient runs background tasks to completion synchronously).
    db = SessionLocal()
    run = PipelineRun(project_id=project["id"], status="RUNNING")
    db.add(run)
    db.commit()
    run_id = run.id
    db.close()

    bus.start_run(run_id)
    bus.request_cancel(run_id)
    agents.runner.run_pipeline(run_id, project["id"], "test-key", "test-model")

    assert client.get(f"/pipeline/{run_id}").json()["status"] == "CANCELLED"

    stream = client.get(f"/pipeline/{run_id}/stream")
    assert "event: error" in stream.text
    assert '"stage": "cancelled"' in stream.text

    # Cancelling an already-finished run is a conflict
    response = client.post(f"/pipeline/{run_id}/cancel")
    assert response.status_code == 409
    assert response.json()["error"]["code"] == "NOT_RUNNING"


def test_sse_replays_events_for_late_connections(client, tmp_path):
    """SSE replay (spec §11): connecting after the run finished must replay the
    stored events and close with the done event."""
    project = client.post("/projects", json={
        "client_name": "Late Co", "project_name": "Replay",
        "project_date": "2026-06-11", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "replay"),
    }).json()
    client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.txt", b"WOs must auto-close.", "text/plain")},
    )
    run_id = client.post(f"/projects/{project['id']}/generate").json()["run_id"]
    assert client.get(f"/pipeline/{run_id}").json()["status"] == "DONE"

    # First late connection replays from the in-memory bus
    stream = client.get(f"/pipeline/{run_id}/stream")
    assert "event: progress" in stream.text
    assert "event: done" in stream.text

    # After a server restart the bus is empty — a final event is synthesized
    bus._events.clear()
    bus._finished.clear()
    stream = client.get(f"/pipeline/{run_id}/stream")
    assert "event: done" in stream.text


def test_pipeline_uses_branded_structure(client, tmp_path):
    """With a branded template uploaded, generation still completes (the branded
    headings replace the default structure in pre-flight, spec §12.1)."""
    import io

    import docx as python_docx

    project = client.post("/projects", json={
        "client_name": "Brand Co", "project_name": "Branded Run",
        "project_date": "2026-06-11", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "branded-run"),
    }).json()
    client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.txt", b"WOs must auto-close.", "text/plain")},
    )

    document = python_docx.Document()
    document.add_heading("Executive Summary", level=1)
    document.add_heading("Business Requirements", level=1)
    buffer = io.BytesIO()
    document.save(buffer)
    branding = client.put(
        f"/projects/{project['id']}/branding",
        files={"file": ("template.docx", buffer.getvalue())},
    )
    assert branding.status_code == 200

    run_id = client.post(f"/projects/{project['id']}/generate").json()["run_id"]
    run = client.get(f"/pipeline/{run_id}").json()
    assert run["status"] == "DONE", run.get("error_message")


def test_upload_uses_embedded_document_date(client, tmp_path):
    """A DOCX downloaded from email keeps its real date inside the file — that
    must win over the browser-reported (reset) modified time."""
    import io
    from datetime import datetime, timezone

    import docx as python_docx

    project = client.post("/projects", json={
        "client_name": "Date Co", "project_name": "Embedded Dates",
        "project_date": "2026-06-12", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "dates"),
    }).json()

    document = python_docx.Document()
    document.add_paragraph("Workshop notes")
    document.core_properties.modified = datetime(2026, 2, 1, 10, 0, tzinfo=timezone.utc)
    buffer = io.BytesIO()
    document.save(buffer)

    upload = client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.docx", buffer.getvalue())},
        data={"source_timestamp": "2026-06-12T08:00:00Z"},  # fake "download time"
    ).json()
    assert upload["source_timestamp"].startswith("2026-02-01T10:00:00")

    # Plain text has no embedded date — the browser-reported time is the fallback
    txt = client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.txt", b"plain notes", "text/plain")},
        data={"source_timestamp": "2026-06-12T08:00:00Z"},
    ).json()
    assert txt["source_timestamp"].startswith("2026-06-12T08:00:00")


def test_refresh_dates_backfills_existing_sources(client, tmp_path):
    """refresh-dates re-reads embedded dates from files already on disk; where
    a date is found it replaces any manual override too."""
    import io
    from datetime import datetime, timezone

    import docx as python_docx

    from db.models import Source

    project = client.post("/projects", json={
        "client_name": "Backfill Co", "project_name": "Refresh Dates",
        "project_date": "2026-06-12", "maximo_version": "mas-9",
        "folder_path": str(tmp_path / "backfill"),
    }).json()

    document = python_docx.Document()
    document.add_paragraph("Old workshop notes")
    document.core_properties.modified = datetime(2026, 2, 1, 10, 0, tzinfo=timezone.utc)
    buffer = io.BytesIO()
    document.save(buffer)
    source_id = client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.docx", buffer.getvalue())},
    ).json()["id"]

    # Simulate a source saved before metadata extraction existed: wrong stored date
    db = SessionLocal()
    db.get(Source, source_id).source_timestamp = datetime(2026, 6, 12, tzinfo=timezone.utc)
    db.commit()
    db.close()

    refreshed = client.post(f"/projects/{project['id']}/sources/refresh-dates").json()
    assert refreshed[0]["source_timestamp"].startswith("2026-02-01T10:00:00")

    # A manual override is wiped when the file yields a real date — the whole
    # point of the button is restoring the files' actual dates
    client.patch(
        f"/projects/{project['id']}/sources/{source_id}",
        json={"user_timestamp_override": "2026-03-05T09:00:00Z"},
    )
    refreshed = client.post(f"/projects/{project['id']}/sources/refresh-dates").json()
    assert refreshed[0]["user_timestamp_override"] is None
    assert refreshed[0]["source_timestamp"].startswith("2026-02-01T10:00:00")

    # A plain-text source has no embedded date, so its override is kept
    txt_id = client.post(
        f"/projects/{project['id']}/sources/upload",
        files={"file": ("notes.txt", b"plain notes", "text/plain")},
    ).json()["id"]
    client.patch(
        f"/projects/{project['id']}/sources/{txt_id}",
        json={"user_timestamp_override": "2026-04-01T12:00:00Z"},
    )
    refreshed = client.post(f"/projects/{project['id']}/sources/refresh-dates").json()
    txt_source = next(s for s in refreshed if s["id"] == txt_id)
    assert txt_source["user_timestamp_override"].startswith("2026-04-01T12:00:00")


def test_validation_errors_use_the_error_envelope(client):
    response = client.post("/projects", json={"client_name": "Only Client"})
    assert response.status_code == 422
    body = response.json()
    assert body["error"]["code"] == "VALIDATION"
    assert "project_name" in body["error"]["message"]
