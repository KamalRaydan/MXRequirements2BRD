"""Processors extract non-empty text from real sample files (spec §17.1)."""
import docx as python_docx
import openpyxl
import pymupdf
import pytest

import processors


def test_classify():
    assert processors.classify("notes.PDF") == "pdf"
    assert processors.classify("req.docx") == "docx"
    assert processors.classify("readme.md") == "plaintext"
    assert processors.classify("data.xlsx") == "spreadsheet"
    assert processors.classify("call.mp3") == "audio"
    assert processors.classify("demo.mov") == "video"
    assert processors.classify("shot.png") == "image"
    assert processors.classify("weird.zzz") == "unknown"


def test_plaintext(tmp_path):
    f = tmp_path / "notes.txt"
    f.write_text("Work orders must auto-close after 30 days.")
    text, pages = processors.extract_text("plaintext", str(f))
    assert "auto-close" in text
    assert pages is None


def test_pdf(tmp_path):
    f = tmp_path / "sample.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "PM schedules are generated weekly.")
    doc.save(str(f))
    doc.close()

    text, pages = processors.extract_text("pdf", str(f))
    assert "PM schedules" in text
    assert pages == 1


def test_docx_paragraphs_and_tables(tmp_path):
    f = tmp_path / "sample.docx"
    document = python_docx.Document()
    document.add_paragraph("Assets must carry criticality ratings.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "REQ-1"
    table.rows[0].cells[1].text = "Inventory reorder automation"
    document.save(str(f))

    text, _ = processors.extract_text("docx", str(f))
    assert "criticality ratings" in text
    assert "REQ-1 | Inventory reorder automation" in text


def test_xlsx(tmp_path):
    f = tmp_path / "reqs.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Requirements"
    ws.append(["ID", "Requirement"])
    ws.append(["R1", "SLA escalation after 4 hours"])
    wb.save(str(f))

    text, _ = processors.extract_text("spreadsheet", str(f))
    assert "Sheet: Requirements" in text
    assert "R1 | SLA escalation after 4 hours" in text


def test_corrupt_pdf_raises(tmp_path):
    f = tmp_path / "broken.pdf"
    f.write_bytes(b"this is not a pdf")
    with pytest.raises(Exception):
        processors.extract_text("pdf", str(f))
