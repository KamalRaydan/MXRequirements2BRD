"""Rendered BRD has requirement tables, DRAFT watermark, named styles (spec §17.1).

Also covers the Milestone 6 branding clone: rendering into a copy of a reference
DOCX carries its fonts, colours, and header logo into the generated BRD.
"""
import json
import struct
import zlib
from io import BytesIO

import docx as python_docx
from docx.shared import Inches, RGBColor

import config
from models.pipeline import BRDDocument, NarrativeSection, Requirement
from services import docx_renderer


def _png_1x1() -> bytes:
    """A real 1x1 RGB PNG with correct CRCs (python-docx's parser is strict)."""
    def chunk(kind: bytes, data: bytes) -> bytes:
        body = kind + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00\x00\x00\x00")
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


_PNG_1X1 = _png_1x1()


def _sample_brd() -> BRDDocument:
    structure = json.loads(
        (config.TEMPLATES_DIR / "brd_default_structure.json").read_text()
    )["sections"]
    return BRDDocument(
        project_metadata={
            "client_name": "Acme Corp",
            "project_name": "EAM Upgrade",
            "project_date": "2026-06-01",
            "maximo_version": "mas-9",
            "maximo_version_label": "MAS 9.x",
        },
        structure=structure,
        narratives=[
            NarrativeSection(section_id="executive_summary", title="Executive Summary",
                             body="First paragraph.\n\nSecond paragraph."),
            NarrativeSection(section_id="risks", title="Implementation Risks",
                             body="Risk one."),
        ],
        requirements=[
            Requirement(id="BRD-WO-001", module="WO", title="Auto-close work orders",
                        description="Close WOs after 30 days", requirement_type="functional",
                        priority="high", source_ref="workshop.pdf",
                        source_timestamp="2026-05-15T09:00:00Z"),
            Requirement(id="BRD-ASSET-001", module="ASSET", title="Criticality ratings",
                        description="All assets rated 1-5", requirement_type="configuration",
                        priority="medium", source_ref="assets.xlsx",
                        source_timestamp="2026-05-16T09:00:00Z"),
        ],
        appendix_sources=["workshop.pdf", "assets.xlsx"],
    )


def test_render_output(tmp_path):
    out = str(tmp_path / "brd.docx")
    docx_renderer.render(_sample_brd(), out)
    document = python_docx.Document(out)

    # One module table per module + the Document Control table
    assert len(document.tables) == 3

    # Requirement IDs present in table cells
    all_cell_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "BRD-WO-001" in all_cell_text
    assert "BRD-ASSET-001" in all_cell_text

    # DRAFT watermark present in the page header XML
    header_xml = document.sections[0].header.paragraphs[0]._p.xml
    assert "DRAFT" in header_xml

    # Headings use named styles; narrative paragraphs use Normal
    styles_used = {p.style.name for p in document.paragraphs}
    assert "Heading 1" in styles_used
    assert "Normal" in styles_used

    # Optional empty sections (non_functional, integrations) are skipped
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert "Non-Functional Requirements" not in headings
    assert "Executive Summary" in headings


def _branded_template(path):
    """Reference DOCX with a distinctive heading font/colour and a header logo."""
    template = python_docx.Document()
    template.styles["Normal"].font.name = "Georgia"
    template.styles["Heading 1"].font.name = "Georgia"
    template.styles["Heading 1"].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    run = template.sections[0].header.paragraphs[0].add_run()
    run.add_picture(BytesIO(_PNG_1X1), width=Inches(1))
    template.save(str(path))
    return str(path)


def test_render_into_branded_template_clones_styles_and_logo(tmp_path):
    template = _branded_template(tmp_path / "reference.docx")
    out = str(tmp_path / "branded-brd.docx")
    profile = {"table_style": "Table Grid"}

    docx_renderer.render(_sample_brd(), out, template_path=template, profile=profile)
    document = python_docx.Document(out)

    # The template's named styles carry over, so the BRD inherits the client's look
    assert document.styles["Heading 1"].font.name == "Georgia"
    assert str(document.styles["Heading 1"].font.color.rgb) == "1F3864"
    assert document.styles["Normal"].font.name == "Georgia"

    # The header logo from the template survives into the output package
    image_parts = [p for p in document.part.package.iter_parts()
                   if p.content_type.startswith("image/")]
    assert image_parts, "expected the branded header logo to be carried over"

    # Our content + watermark still render correctly inside the branded shell
    assert len(document.tables) == 3
    assert "DRAFT" in document.sections[0].header.paragraphs[0]._p.xml
    assert "Executive Summary" in [
        p.text for p in document.paragraphs if p.style.name.startswith("Heading")
    ]


def test_apply_table_style_falls_back_for_unknown_style(tmp_path):
    out = str(tmp_path / "fallback.docx")
    # An unknown table style must not crash rendering — it falls back to Table Grid
    docx_renderer.render(_sample_brd(), out, profile={"table_style": "No Such Style"})
    document = python_docx.Document(out)
    assert document.tables[0].style.name == "Table Grid"
