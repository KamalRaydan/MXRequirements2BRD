"""Read the visual identity ("branding") of a client's reference DOCX (Milestone 6).

The actual visual clone happens by rendering the BRD *into a copy of the reference
DOCX* (see docx_renderer) so its fonts, colours, theme, header logo, and table
styles carry over natively. This module extracts a small, human-readable summary
of what that document contains — body/heading fonts, the heading colour, the table
style, and whether a logo was found — so the UI can confirm what was detected and
the renderer knows which table style to reuse.

Everything here is best-effort: Word documents leave fonts/colours unset when they
inherit them from the theme, so any field may legitimately be None.
"""
import json
from pathlib import Path

from docx import Document

# Where the per-project profile + extracted logo live (next to reference.docx).
PROFILE_FILENAME = "profile.json"
LOGO_STEM = "logo"


def _style_font_name(document, style_name: str) -> str | None:
    """Font name set directly on a named style, or None if it inherits the theme."""
    try:
        return document.styles[style_name].font.name
    except KeyError:
        return None


def _style_font_size_pt(document, style_name: str) -> float | None:
    try:
        size = document.styles[style_name].font.size
    except KeyError:
        return None
    return round(size.pt, 1) if size is not None else None


def _style_color_hex(document, style_name: str) -> str | None:
    """Explicit RGB colour of a named style as a hex string, or None.

    Theme colours (which have no fixed RGB) and unset colours both return None.
    """
    try:
        color = document.styles[style_name].font.color
    except KeyError:
        return None
    try:
        rgb = color.rgb  # raises / is None for theme or unset colours
    except Exception:
        return None
    return str(rgb) if rgb is not None else None


def _first_table_style(document) -> str | None:
    if not document.tables:
        return None
    try:
        return document.tables[0].style.name
    except Exception:
        return None


def _find_logo_part(document):
    """First embedded image in the document (header logos included).

    Returns the python-docx Part, or None. Embedded images live under /word/media/;
    we deliberately ignore /docProps/thumbnail.jpeg (an auto-generated preview that
    Word stores in every document, not part of the branding). We don't try to guess
    which image is "the logo" — for a branding template the first one is it.
    """
    for part in document.part.package.iter_parts():
        if part.content_type.startswith("image/") and str(part.partname).startswith("/word/media/"):
            return part
    return None


def _save_logo(part, branding_dir: Path) -> str | None:
    """Write the logo image into branding/ and return its path."""
    # Reuse the original extension so Word/preview tools recognise it.
    suffix = Path(part.partname).suffix or ".png"
    logo_path = branding_dir / f"{LOGO_STEM}{suffix}"
    logo_path.write_bytes(part.blob)
    return str(logo_path)


def extract_profile(reference_path: str, branding_dir: str | None = None) -> dict:
    """Summarise the reference DOCX's visual branding.

    When `branding_dir` is given, the logo image (if any) is saved there and its
    path is included; otherwise only `logo_found` is reported (the pipeline gets
    the real logo for free from the template, so it doesn't need the file).
    """
    document = Document(reference_path)
    logo_part = _find_logo_part(document)

    profile = {
        "body_font": _style_font_name(document, "Normal"),
        "body_size_pt": _style_font_size_pt(document, "Normal"),
        "heading_font": _style_font_name(document, "Heading 1"),
        "heading_color": _style_color_hex(document, "Heading 1"),
        "table_style": _first_table_style(document),
        "logo_found": logo_part is not None,
        "logo_path": None,
    }

    if branding_dir and logo_part is not None:
        profile["logo_path"] = _save_logo(logo_part, Path(branding_dir))

    return profile


def save_profile(profile: dict, branding_dir: str) -> None:
    path = Path(branding_dir) / PROFILE_FILENAME
    path.write_text(json.dumps(profile, indent=2), encoding="utf-8")


def load_profile(branding_dir: str) -> dict | None:
    path = Path(branding_dir) / PROFILE_FILENAME
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


def remove_profile(branding_dir: str) -> None:
    """Delete the stored profile.json and any logo image."""
    directory = Path(branding_dir)
    (directory / PROFILE_FILENAME).unlink(missing_ok=True)
    for logo in directory.glob(f"{LOGO_STEM}.*"):
        logo.unlink(missing_ok=True)
