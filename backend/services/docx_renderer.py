"""Render a BRDDocument to .docx (spec §12.6).

Rules: named Word styles only (Heading 1/2/3, Normal, List Bullet, Table Grid),
requirements as one table per module, DRAFT watermark injected into the page
header as WordArt XML (so it appears behind every page, not as a text box).

Branding clone (Milestone 6): when a project has a branded reference DOCX we render
*into a cleared copy of that document* instead of a blank one. Word keeps fonts,
colours, theme, and the header logo as named styles / separate parts, so adding our
headings and paragraphs with the same named styles makes them inherit the client's
look automatically — a high-fidelity clone with no fragile per-run XML surgery. The
detected table style (from the BrandingProfile) is reused for our tables.
"""
from pathlib import Path

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches

from models.pipeline import BRDDocument, Requirement

# Usable page width on US Letter with 1" margins
_PAGE_WIDTH_INCHES = 6.5

# Column widths as ratios of page width (spec §9.3)
_REQ_COLUMNS = [
    ("Requirement ID", 0.15),
    ("Title", 0.20),
    ("Description", 0.35),
    ("Type", 0.10),
    ("Priority", 0.10),
    ("Source Document", 0.10),
]

# WordArt watermark, same XML Word itself emits for an inserted watermark.
# The shapetype defines the WordArt geometry; the shape is the rotated grey "DRAFT".
_WATERMARK_XML = (
    '<w:pict {w} xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office">'
    '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" '
    'path="m@7,l@8,m@5,21600l@6,21600e">'
    '<v:formulas>'
    '<v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>'
    '<v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/>'
    '<v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/>'
    '<v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/>'
    '<v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/>'
    '<v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/>'
    '<v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/>'
    '</v:formulas>'
    '<v:path textpathok="t" o:connecttype="custom" '
    'o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>'
    '<v:textpath on="t" fitshape="t"/>'
    '<v:handles><v:h position="#0,bottomRight" xrange="6629,14971"/></v:handles>'
    '<o:lock v:ext="edit" text="t" shapetype="t"/>'
    '</v:shapetype>'
    '<v:shape id="MaximoBRDWatermark" type="#_x0000_t136" '
    'style="position:absolute;margin-left:0;margin-top:0;width:468pt;height:175.5pt;'
    'rotation:315;z-index:-251654144;mso-position-horizontal:center;'
    'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
    'mso-position-vertical-relative:margin" o:allowincell="f" fillcolor="silver" stroked="f">'
    '<v:fill opacity=".5"/>'
    '<v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt" string="DRAFT"/>'
    '</v:shape>'
    '</w:pict>'
).format(w=nsdecls("w"))


def _add_draft_watermark(document: Document) -> None:
    # Append the watermark run to the header's first paragraph. In branded mode the
    # header may already hold the client's logo; the watermark floats (absolute
    # position) so it overlays without displacing anything already there.
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run._r.append(parse_xml(_WATERMARK_XML))


def _open_document(template_path: str | None) -> Document:
    """Blank document, or a cleared copy of the branded template (Milestone 6).

    Opening the reference DOCX inherits all its named styles, theme, and header
    (logo); we only strip the client's body content, keeping the trailing section
    properties that point at those headers and define the page layout.
    """
    if template_path and Path(template_path).exists():
        document = Document(template_path)
        _clear_body(document)
        return document
    return Document()


def _clear_body(document: Document) -> None:
    """Remove the template's body paragraphs and tables, keep <w:sectPr>.

    Styles, theme, and headers/footers live in separate package parts and are
    untouched; the final body-level sectPr carries page size/margins + header refs.
    """
    body = document.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _apply_table_style(table, table_style: str) -> None:
    """Use the detected/branded table style, falling back to Table Grid if the
    document doesn't define it (best-effort — spec §18)."""
    for candidate in (table_style, "Table Grid"):
        try:
            table.style = candidate
            return
        except (KeyError, ValueError):
            continue


def _add_document_control(document: Document, metadata: dict, table_style: str) -> None:
    table = document.add_table(rows=0, cols=2)
    _apply_table_style(table, table_style)
    rows = [
        ("Client", metadata.get("client_name", "")),
        ("Project", metadata.get("project_name", "")),
        ("Project Date", metadata.get("project_date", "")),
        ("Maximo Version", metadata.get("maximo_version_label", metadata.get("maximo_version", ""))),
        ("Document Status", "DRAFT"),
    ]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].text = str(value)


def _add_requirements_table(document: Document, requirements: list[Requirement],
                            table_style: str) -> None:
    table = document.add_table(rows=1, cols=len(_REQ_COLUMNS))
    _apply_table_style(table, table_style)

    for i, (title, ratio) in enumerate(_REQ_COLUMNS):
        table.columns[i].width = Inches(_PAGE_WIDTH_INCHES * ratio)
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].add_run(title).bold = True

    for req in requirements:
        cells = table.add_row().cells
        values = [req.id, req.title, req.description, req.requirement_type,
                  req.priority, req.source_ref]
        for i, value in enumerate(values):
            cells[i].text = value


def render(brd: BRDDocument, output_path: str,
           template_path: str | None = None, profile: dict | None = None) -> str:
    document = _open_document(template_path)
    _add_draft_watermark(document)

    # Reuse the branded document's table style when we detected one (Milestone 6).
    table_style = (profile or {}).get("table_style") or "Table Grid"

    narratives_by_id = {n.section_id: n for n in brd.narratives}
    requirements_by_module: dict[str, list[Requirement]] = {}
    for req in brd.requirements:
        requirements_by_module.setdefault(req.module, []).append(req)

    for section in brd.structure:
        section_id = section["id"]
        narrative = narratives_by_id.get(section_id)
        is_special = section_id in ("cover", "requirements", "appendix")

        # Skip optional sections that ended up with no content
        if not is_special and not narrative and not section.get("required", False):
            continue

        document.add_heading(section["title"], level=min(section.get("level", 1), 3))

        if section_id == "cover":
            _add_document_control(document, brd.project_metadata, table_style)
        elif section_id == "requirements":
            for module in sorted(requirements_by_module):
                document.add_heading(module, level=2)
                _add_requirements_table(document, requirements_by_module[module], table_style)
        elif section_id == "appendix":
            for filename in brd.appendix_sources:
                document.add_paragraph(filename, style="List Bullet")
        elif narrative:
            for paragraph_text in narrative.body.split("\n\n"):
                if paragraph_text.strip():
                    document.add_paragraph(paragraph_text.strip())

    document.save(output_path)
    return output_path
